import tkinter
import customtkinter
import openai
import uuid
from docx import Document
from docx.shared import Pt
from docx2pdf import convert
import os
from tkinter import filedialog
from tkinter import *
import subprocess


# declarations
questions = []
randomFileName = ''


# Define OpenAI API key 
openai.api_key = ""
model_engine = "text-davinci-003"



# functions
def saveFileLocation():
    app.filename =  filedialog.askopenfilename(initialdir = "/",title = "Select file",filetypes = (("Word document","*.docx"),("all files","*.*")))
    fileNameLabel.configure(text="Selected File : " + app.filename, text_color="green")


def readDocx():
    f = open(app.filename, 'rb')
    doc = Document(f)
    f.close()
    questions.clear()
    for docpara in doc.paragraphs:
        questions.append(docpara.text)

def solveFunction():
    try:
            readDocx()
            randomFileName = "answers" + str(uuid.uuid4())
            document = Document()
            document.add_heading('Solutions', 0)
            for x in range(len(questions)):
                prompt = questions[x]
                completion = openai.Completion.create(
                engine=model_engine,
                prompt=prompt,
                max_tokens=1024,
                n=1,
                stop=None,
                temperature=0.5,
                )
                response = completion.choices[0].text
    
                quest = document.add_paragraph('', style='List Number').add_run(questions[x].strip())
                ans = document.add_paragraph().add_run(response.strip())
                quest.font.size = Pt(12)
                quest.bold = True
                ans.font.size = Pt(12)
            
            solveLabel.configure(text="Saving PDF", text_color="green")
            document.save(randomFileName + ".docx")
            convert(randomFileName + ".docx")
            os.remove(randomFileName + ".docx")
            solveLabel.configure(text="PDF Saved Successfully", text_color="green")
            subprocess.Popen([randomFileName + ".pdf"],shell=True)
    except:
         solveLabel.configure(text="Please Try Again Later", text_color="red")


# GUI Elements
app = customtkinter.CTk()
app.geometry("720x480")
app.title("Question Bank Solver")
customtkinter.set_appearance_mode("System")
customtkinter.set_default_color_theme("blue")

browseButton = customtkinter.CTkButton(app, text="Select Question Bank (Word Document)", command=saveFileLocation)
browseButton.pack(padx=10, pady=10)

fileNameLabel = customtkinter.CTkLabel(app, text="")
fileNameLabel.pack()

solveButton = customtkinter.CTkButton(app, text="Solve and Save PDF", command=solveFunction)
solveButton.pack(padx=10, pady=10)

solveLabel = customtkinter.CTkLabel(app, text="")
solveLabel.pack()

app.mainloop()